import sys
import os
import tkinter as tk
from PIL import ImageTk, Image
from openpyxl import load_workbook
from tkinter import messagebox, ttk
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_COLOR_INDEX


# ONE FILE TEMP FILE
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

constan_path = resource_path("Constanta.xlsx")

# def Agar input menjadi angka saja 
def validate_input(new_value):
    if new_value.isdigit() or new_value == "." or new_value == "9.81":
        return True
    else:
        return False
# def entry max 50
def validate_50(text):
    if text.isdigit():
        if int(text) > 50:
            return False
        else:
            return True
    elif text == "":
        return True
    else:
        return False



"""Code Aplikasinya"""
# class aplikasinya
class App(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.master.title("Daya Dukung Pada Pondasi Dangkal")
        self.create_widgets()


    def create_widgets(self):
        ilusap_path = resource_path("ilus1.png")
        logo_path = resource_path("logo.png")
        self.ilus = ImageTk.PhotoImage(Image.open(ilusap_path).resize((211, 141), Image.BILINEAR))
        tk.Label(self.master, image=self.ilus).place(x=380, y=150)
        self.logoits = ImageTk.PhotoImage(Image.open(logo_path).resize((140, 50), Image.BILINEAR))
        tk.Label(self.master, image=self.logoits).place(x=550, y=10)
        

        """Combobox Code"""
        # create a list of options
        self.options = ["General Shear", "Local Shear"]

        # create a variable to store the selected option
        self.selected_option = tk.StringVar()

        # create a combobox and add options to it
        tk.Label(self.master, text="Pilih Tipe Shear :").grid(row=1, column=4, sticky="w")
        self.combobox = ttk.Combobox(self.master, textvariable=self.selected_option, values=self.options, state="readonly")
        self.combobox.grid(row=2, column=4)

        # set the default option
        self.selected_option.set(self.options[0])

        def handle_selection(event):
            self.selected = self.selected_option.get()
            print("Selected option:", self.selected)

        # bind the combobox to the selection function
        self.combobox.bind("<<ComboboxSelected>>", handle_selection)

        """Code Widged lainnya"""

        tk.Label(self.master, text="  ",font=("Arial", 12)).grid(row=0, column=1, sticky="w")
        tk.Label(self.master, text="             ").grid(row=1, column=0, sticky="w")

        # Input Box
        tk.Label(self.master, text="Berat jenis tanah γ").grid(row=1, column=1, sticky="w")
        self.gamma_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6,font=("Arial", 12))
        self.gamma_entry.grid(row=1, column=2)
        tk.Label(self.master, text="(kN/m3)").grid(row=1, column=3, padx=0, pady=0, sticky="w")

        tk.Label(self.master, text="Sudut gesekan dalam Φ'").grid(row=2, column=1, sticky="w")
        self.phi_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_50), '%P'),width=6,font=("Arial", 12))
        self.phi_entry.grid(row=2, column=2)
        tk.Label(self.master, text="° (max = 50)             ").grid(row=2, column=3, pady=0, sticky="w")

        tk.Label(self.master, text="Kohesi tanah c'").grid(row=3, column=1, sticky="w")
        self.c_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6,font=("Arial", 12))
        self.c_entry.grid(row=3, column=2)
        tk.Label(self.master, text="(kN/m3)").grid(row=3, column=3, padx=0, pady=0, sticky="w")


        tk.Label(self.master, text="Lebar pondasi B").grid(row=4, column=1, sticky="w")
        self.B_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6,font=("Arial", 12))
        self.B_entry.grid(row=4, column=2)
        tk.Label(self.master, text="(m)").grid(row=4, column=3, padx=0, pady=0, sticky="w")


        tk.Label(self.master, text="Kedalaman pondasi Df").grid(row=5, column=1, sticky="w")
        self.df_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6, font=("Arial", 12))
        self.df_entry.grid(row=5, column=2)
        tk.Label(self.master, text="(m)").grid(row=5, column=3, padx=0, pady=0, sticky="w")


        tk.Label(self.master, text="Muka air tanah MAT").grid(row=6, column=1, sticky="w")
        self.mat_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6, font=("Arial", 12))
        self.mat_entry.grid(row=6, column=2)
        tk.Label(self.master, text="(m) dibawah tanah      ").grid(row=6, column=3, pady=0, sticky="w")


        tk.Label(self.master, text="Berat jenis air γw").grid(row=7, column=1, sticky="w")
        self.yw_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6, font=("Arial", 12))
        self.yw_entry.grid(row=7, column=2)
        self.yw_entry.insert(0, "9.81")
        tk.Label(self.master, text="(kN/m3)").grid(row=7, column=3, padx=0, pady=0, sticky="w")


        tk.Label(self.master, text="Berat jenis tanah jenuh γsat").grid(row=8, column=1, sticky="w")
        self.ysat_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6, font=("Arial", 12))
        self.ysat_entry.grid(row=8, column=2)
        tk.Label(self.master, text="(kN/m3)").grid(row=8, column=3, padx=0, pady=0, sticky="w")


        tk.Label(self.master, text="Angka keamanan SF").grid(row=9, column=1, sticky="w")
        self.sf_entry = tk.Spinbox(self.master, from_=1, to_=20, increment=0.1, textvariable=tk.IntVar(value=3),width=5 ,font=("Arial", 12))
        self.sf_entry.grid(row=9, column=2,sticky=tk.W)

        # boolean checkbox
        self.check_1 = tk.BooleanVar(value=True)
        self.check_2 = tk.BooleanVar(value=True)
        self.check_3 = tk.BooleanVar(value=True)

        # membuat checkbox
        self.cek1 = tk.Checkbutton(self.master, text="Pondasi Menerus", variable=self.check_1).grid(row=3, column=4, padx=0, pady=0, sticky=tk.W)
        self.cek2 = tk.Checkbutton(self.master, text="Pondasi Bujursangkar", variable=self.check_2).grid(row=4, column=4, sticky=tk.W)
        self.cek3 = tk.Checkbutton(self.master, text="Pondasi Lingkaran", variable=self.check_3).grid(row=5, column=4, sticky=tk.W)

        self.selected = self.selected_option.get()

        # Calculate Button
        tk.Button(self.master, text="Hitung Daya Dukung", command=self.calculate).grid(row=10, column=2, rowspan=2, columnspan=2,  padx=0, pady=25)
        # Info Author Button
        tk.Button(self.master, text="Info", command=self.infobruh).grid(row=11, column=1,  padx=0, pady=15, sticky=tk.W)

    def infobruh(self):
        tk.messagebox.showinfo("Info", "Program ini dibuat oleh :\n\n   1. Ridho Rizky Febriansyah 2035221071 (Coder)")#\n   2. Izzati Eka Nisa 2035221047 (Team)\n   3. Nadia Aura Deka 2035221050 (Team)\n   4. Fiza Eka Salsabilah 2035221084 (Team)\n   5. Salsabilla shafi A 2035221089 (Team)\n   6. Affarel Sebatian N 2035221056 (Team)\n   7. Abdillah Pandega W. 2035221065 (Team)\n   8. Samudro Luhur Trapsilo 2035221086 (Team)\n   9. Syarifuddin Putra P 2035221053 (Team)\n   10. I Made Raghayana 2035221057 (Team)\n   11. Barka Hilman P. 2035221051 (Team)\n")


    def calculate(self):
        # Get input values
        try:
            gamma = float(self.gamma_entry.get())
            phi = float(self.phi_entry.get())
            c = float(self.c_entry.get())
            B = float(self.B_entry.get())
            df = float(self.df_entry.get())
            sf = float(self.sf_entry.get())
            mat = float(self.mat_entry.get())
            yw = float(self.yw_entry.get())
            ysat = float(self.ysat_entry.get())
        # Jendela eror saat vaule salah
        except Exception as e:
            messagebox.showerror("Error", "Tolong Input data dengan benar")


        self.cekcekbox_ya = self.check_1.get()+self.check_2.get()+self.check_3.get()
        
        if self.cekcekbox_ya == 0:
            messagebox.showerror("Pilih salah satu bruhh", "Tolong Setidaknya Pilih salah satu pilihan jenis pondasi")
            B = str(c)

        # Memilih data shear
        if self.selected == "General Shear":
            rckor = 0
            qckonbl = 1.3
            qckonc = 1
        else:
            rckor = 4
            qckonbl = 0.867
            qckonc = 2/3

        d = abs(mat-df)

        colphi = 3 + phi
        wblod = load_workbook(constan_path)
        ws1 = wblod.active  
        Nc = ws1.cell(row=colphi, column=2+rckor ).value
        Nq = ws1.cell(row=colphi, column=3+rckor ).value
        Ny = ws1.cell(row=colphi, column=4+rckor ).value

        # perhitungan
        gammaaksen = ysat - yw
        yav = ((gamma*(mat-df)+gammaaksen*(B-(mat-df)))/B)

        # Pengaruh MAT Kondisi 1 2 dan 3
        if df > mat:
            q = gamma * mat + gammaaksen * (df-mat)
            gamma = gammaaksen
            cek_mat = "1"
            rumus_ikut_mat = 1
            kondisi_mat = "di atas"
        elif df == mat:
            q = gamma * df
            gamma = gammaaksen
            cek_mat = "2"
            rumus_ikut_mat = 2
            kondisi_mat = "pas di"
        elif df < mat:
            cek_mat = "3"
            rumus_ikut_mat = 3
            kondisi_mat = "di bawah"
            if mat-df < B:
                gamma = yav
                q = yav * df
                rumus_ikut_mat_kon = 1
            elif mat-df == B:
                gamma = yav
                q = yav * df
                rumus_ikut_mat_kon = 1
            elif mat-df > B:
                q = df * gamma
                rumus_ikut_mat_kon = 2

        qq = Nq * q


        # Pondasi menerus
        if self.check_1.get() == 1:
            qc = Nc * c * qckonc
            qp = Ny * B * (gamma / 2)
            qu = (qc + qq + qp)
            qall = (qu / sf) 
        # Pondasi Bujursangkar
        if self.check_2.get() == 1:
            qca2 = Nc * c * qckonbl
            qpa2 = Ny * B * gamma * 0.4
            qua2 = (qca2 + qq + qpa2)
            qall2 = (qua2 / sf)
        # Pondasi lingkaran
        if self.check_3.get() == 1:
            qca3 = Nc * c * qckonbl
            qpa3 = Ny * B * (gamma * 0.3)
            qua3 = (qca3 + qq + qpa3)
            qalla3 = (qua3 / sf)

        
        # Simpaan Hasil ke dalam file Word
        # membuat objek document
        gamma = float(self.gamma_entry.get())
        doc = docx.Document()
        # Add a Title to the document 
        doc.add_heading('Daya Dukung Pondasi Dangkal', 0)

        # menambahkan teks ke dokumen
        para = doc.add_paragraph('Diketahui : ').add_run()
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        table = doc.add_table(rows=2, cols=9)

        # mengisi sel-sel tabel dengan teks
        table.cell(0, 0).text = "γ"
        table.cell(0, 1).text = "γsat"
        table.cell(0, 2).text = "γw"
        table.cell(0, 3).text = "Φ'"
        table.cell(0, 4).text = "c'"
        table.cell(0, 5).text = "B"
        table.cell(0, 6).text = "Df"
        table.cell(0, 7).text = "MAT"
        table.cell(0, 8).text = "SF"
        table.cell(1, 0).text = str(gamma)+' (kN/m3)'
        table.cell(1, 1).text = str(ysat)+' (kN/m3)'
        table.cell(1, 2).text = str(yw)+' (kN/m3)'
        table.cell(1, 3).text = str(phi)+'°'
        table.cell(1, 4).text = str(c)+' (kN/m3)'
        table.cell(1, 5).text = str(B)+' (m)'
        table.cell(1, 6).text = str(df)+' (m)'
        table.cell(1, 7).text = str(mat)+' (m)'
        table.cell(1, 8).text = str(sf)

        # menambahkan style tabel baru
        new_style = doc.styles.add_style('MyTableStyle', docx.enum.style.WD_STYLE_TYPE.TABLE)
        new_style.base_style = doc.styles['Table Grid']  # gaya dasar tabel
        new_style.font.name = 'Times New Roman'  # jenis huruf
        new_style.font.size = docx.shared.Pt(12)  # ukuran huruf
        new_style.paragraph_format.alignment = docx.enum.text.WD_TAB_ALIGNMENT.CENTER  # alignment teks dalam sel
        # menerapkan gaya baru pada tabel
        table.style = 'MyTableStyle'


        para = doc.add_paragraph().add_run('\nDitanya :')
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        para = doc.add_paragraph().add_run('Hitung daya dukung pondasi dangkal Formula Therzhagi dengan '+self.selected+'.')
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        # Adding points to the list named 'List Number'
        if self.check_1.get() == 1:
            doc.add_paragraph('Pondasi Menerus.',  style='List Number 2')
        if self.check_2.get() == 1:
            doc.add_paragraph('Pondasi Bujur-Sangkar.',  style='List Number 2')
        if self.check_3.get() == 1:
            doc.add_paragraph('Pondasi Lingkaran.',  style='List Number 2')

        para = doc.add_paragraph().add_run('Dijawab :')
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        # menyisipkan gambar ke dalam dokumen
        img_p = resource_path("ilus"+cek_mat+".png")
        doc.add_picture(img_p)#, width=Cm(2), height=Cm(2))


        para = doc.add_paragraph().add_run("Pada Kondisi tersebut muka air tanah berada "+kondisi_mat+" Dasar pondasi. maka ")
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        dogi = doc.add_paragraph()

        para = doc.add_paragraph()
        para.add_run("D	= |MAT-Df|\n")
        para.add_run(" 	= |"+str(mat)+"-"+str(df)+"|\n")
        para.add_run(" 	= "+str(d)+" m\n")
        para.add_run("γ'	= γsat - γw\n")
        para.add_run("	= "+str(ysat)+" - "+str(yw)+"\n")
        para.add_run("	= "+str(gammaaksen)+" kN/m3\n")

        if rumus_ikut_mat == 2:
            para.add_run("q	= Df x γ\n")
            para.add_run("	= "+str(df)+" x "+str(gamma)+"\n")
            para.add_run("	= "+str(q)+" kN/m2\n")
            gamma = gammaaksen
        elif rumus_ikut_mat == 1:
            para.add_run("q	= γ(MAT) + γ'D\n")
            para.add_run("	= "+str(gamma)+" x "+str(mat)+" + "+str(gammaaksen)+" x "+str(d) +"\n")
            para.add_run("	= "+str(q)+" kN/m2\n")
            gamma = gammaaksen
        else:
            if rumus_ikut_mat_kon == 1:
                para.add_run("γav	= (γ x D + γ' x (B-D)) / B\n")
                para.add_run("	= "+str(gamma)+" x "+str(d)+" + "+str(gammaaksen)+" x ("+str(B)+" - "+str(d)+")) / "+str(B)+"\n")
                para.add_run("	= "+ str(yav) +" kN/m3\n")
                para.add_run("q	= γav x Df\n")
                para.add_run("	= "+str(yav)+" x "+str(df)+"\n")
                para.add_run("	= "+str(q)+" kN/m2\n")
                gamma = yav
            else:
                para.add_run("γav	= γ\n")
                para.add_run("γav	= "+str(gamma)+"\n")
                para.add_run("q	= γ x Df\n")
                para.add_run("	= "+str(gamma)+" x "+str(df)+"\n")
                para.add_run("	= "+str(q)+" kN/m2\n")

        
        

        if rumus_ikut_mat == 1:
            dogi.add_run("Jika Df > MAT maka\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            dogi.add_run("	")
            dogi.add_run("q = γ(MAT) + γ'D\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            dogi.add_run("	")
            dogi.add_run("γ pada suku ke 3 diganti γ'\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif rumus_ikut_mat == 2:
            dogi.add_run("Jika Df = MAT maka\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            dogi.add_run("	q = γ * Df\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            dogi.add_run("	γ pada suku ke 3 diganti γ'\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            dogi.add_run("Jika Df < MAT ").font.highlight_color = WD_COLOR_INDEX.YELLOW
            if rumus_ikut_mat_kon == 1:
                dogi.add_run("dan D ≤ B maka\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
                dogi.add_run("	")
                dogi.add_run("γ diganti γav\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
                dogi.add_run("	")
                dogi.add_run("q = γav * Df\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                dogi.add_run("dan D > B maka\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
                dogi.add_run("	")
                dogi.add_run("γav = γ\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
                dogi.add_run("	")
                dogi.add_run("q = γ * Df\n").font.highlight_color = WD_COLOR_INDEX.YELLOW

        
        if self.selected == "General Shear":
            """General Shear"""
            #menerus
            if self.check_1.get() == 1:
                doc.add_paragraph("Pondasi Menerus",  style='List Number')
                doc.add_paragraph().add_run("qu	= c'(Nc) + q(Nq) + 0.5γB(Nγ)")
                doc.add_paragraph().add_run("	= "+str(c)+" · "+str(Nc)+" + "+str(q)+" · "+str(Nq)+" + 0.5 · "+str(gamma)+" · "+str(B)+" · "+str(Ny))
                doc.add_paragraph().add_run("	= "+str(qc)+" + "+str(qq)+" + "+str(qp))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qu)+" kN/m2").font.highlight_color = WD_COLOR_INDEX.YELLOW

                doc.add_paragraph().add_run("qall	= qu / SF")
                doc.add_paragraph().add_run("	= "+str(qu)+" / "+str(sf))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qall)+" kN/m2\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            #Bujur-Sangkar
            if self.check_2.get() == 1:
                doc.add_paragraph("Pondasi Bujur-Sangkar",  style='List Number')
                doc.add_paragraph().add_run("qu	= 1.3c'(Nc) + q(Nq) + 0.4γB(Nγ)")
                doc.add_paragraph().add_run("	= 1.3 · "+str(c)+" · "+str(Nc)+" + "+str(q)+" · "+str(Nq)+" + 0.4 · "+str(gamma)+" · "+str(B)+" · "+str(Ny))
                doc.add_paragraph().add_run("	= "+str(qca2)+" + "+str(qq)+" + "+str(qpa2))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qua2)+" kN/m2").font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_paragraph().add_run("qall	= qu / SF")
                doc.add_paragraph().add_run("	= "+str(qua2)+" / "+str(sf))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qall2)+" kN/m2\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            # Lingkaran
            if self.check_3.get() == 1:
                doc.add_paragraph("Pondasi Lingkaran",  style='List Number')
                doc.add_paragraph().add_run("qu	= 1.3c'(Nc) + q(Nq) + 0.3γB(Nγ)")
                doc.add_paragraph().add_run("	= 1.3 · "+str(c)+" · "+str(Nc)+" + "+str(q)+" · "+str(Nq)+" + 0.3 · "+str(gamma)+" · "+str(B)+" · "+str(Ny))
                doc.add_paragraph().add_run("	= "+str(qca3)+" + "+str(qq)+" + "+str(qpa3))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qua3)+" kN/m2").font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_paragraph().add_run("qall	= qu / SF")
                doc.add_paragraph().add_run("	= "+str(qua3)+" / "+str(sf))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qalla3)+" kN/m2\n").font.highlight_color = WD_COLOR_INDEX.YELLOW

        else:
            """Local Shear"""
            if self.check_1.get() == 1:
                doc.add_paragraph("Pondasi Menerus",  style='List Number')
                doc.add_paragraph().add_run("qu	= ⅔c'(N'c) + q(N'q) + 0.5γB(N'γ)")
                doc.add_paragraph().add_run("	= ⅔·"+str(c)+" · "+str(Nc)+" + "+str(q)+" · "+str(Nq)+" + 0.5 · "+str(gamma)+" · "+str(B)+" · "+str(Ny))
                doc.add_paragraph().add_run("	= "+str(qc)+" + "+str(qq)+" + "+str(qp))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qu)+" kN/m2").font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_paragraph().add_run("qall	= qu / SF")
                doc.add_paragraph().add_run("	= "+str(qu)+" / "+str(sf))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qall)+" kN/m2\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            #Bujur-Sangkar
            if self.check_2.get() == 1:
                doc.add_paragraph("Pondasi Bujur-Sangkar",  style='List Number')
                doc.add_paragraph().add_run("qu	= 0.867c'(N'c) + q(N'q) + 0.4γB(N'γ)")
                doc.add_paragraph().add_run("	= 0.867·"+str(c)+" · "+str(Nc)+" + "+str(q)+" · "+str(Nq)+" + 0.4 · "+str(gamma)+" · "+str(B)+" · "+str(Ny))
                doc.add_paragraph().add_run("	= "+str(qca2)+" + "+str(qq)+" + "+str(qpa2))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qua2)+" kN/m2").font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_paragraph().add_run("qall	= qu / SF")
                doc.add_paragraph().add_run("	= "+str(qua2)+" / "+str(sf))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qall2)+" kN/m2\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            # Lingkaran
            if self.check_3.get() == 1:
                doc.add_paragraph("Pondasi Lingkaran",  style='List Number')
                doc.add_paragraph().add_run("qu	= 0.867c'(N'c) + q(N'q) + 0.3γB(N'γ)")
                doc.add_paragraph().add_run("	= 0.867·"+str(c)+" · "+str(Nc)+" + "+str(q)+" · "+str(Nq)+" + 0.3 · "+str(gamma)+" · "+str(B)+" · "+str(Ny))
                doc.add_paragraph().add_run("	= "+str(qca3)+" + "+str(qq)+" + "+str(qpa3))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qua3)+" kN/m2").font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_paragraph().add_run("qall	= qu / SF")
                doc.add_paragraph().add_run("	= "+str(qua3)+" / "+str(sf))
                par = doc.add_paragraph()
                par.add_run("	= ")
                par.add_run(str(qalla3)+" kN/m2\n").font.highlight_color = WD_COLOR_INDEX.YELLOW



        # mengatur lebar kolom pada tabel
        kolom = table.columns
        for i in range(len(kolom)):
            kolom[i].width = Cm(1)

        # mengatur lebar sel pada tabel
        sel = table.cell(0, 2)
        sel.width = Cm(1)

        # menyimpan dokumen ke dalam file
        doc.save('Hasil_Perhitungan_DD_Pondasi_Dangkal.docx')


        # Tampilkan pesan berhasil
        tk.messagebox.showinfo("Sukses", "Perhitungan berhasil. Hasil perhitungan tersimpan di Hasil_Perhitungan_DD_Pondasi_Dangkal.docx")

ico_path = resource_path("ico.ico")

if __name__ == "__main__":
    root = tk.Tk()
    root.iconbitmap(ico_path)
    root.geometry("640x330")
    root.resizable(False, False)
    app = App(master=root)
    app.mainloop()

