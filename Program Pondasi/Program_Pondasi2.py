import sys
import os
import tkinter as tk
from PIL import ImageTk, Image
from tkinter import messagebox, ttk
import docx
from docx.shared import Pt

# ONE FILE TEMP FILE
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

ico_path = resource_path("ico.ico")
ilut1_path = resource_path("ilusteg1.png")
ilut2_path = resource_path("ilusteg2.png")
logo_path = resource_path("logo.png")

# def Agar input menjadi angka saja 
def validate_input(new_value):
    if new_value.isdigit() or new_value == ".":
        return True
    else:
        return False


"""Code Aplikasinya"""
# class aplikasinya
class App(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.master.title("Tegangan Di Dasar Pondasi Dangkal")
        self.create_widgets()


    def create_widgets(self):

        
        self.ilus = ImageTk.PhotoImage(Image.open(ilut1_path).resize((180, 141), Image.BILINEAR))
        tk.Label(self.master, image=self.ilus).place(x=280, y=10)
        self.ilus2 = ImageTk.PhotoImage(Image.open(ilut2_path).resize((101, 108), Image.BILINEAR))
        tk.Label(self.master, image=self.ilus2).place(x=470, y=70)
        self.logoits = ImageTk.PhotoImage(Image.open(logo_path).resize((140, 50), Image.BILINEAR))
        tk.Label(self.master, image=self.logoits).place(x=480, y=10)


        """Code Widged lainnya"""


        tk.Label(self.master, text="  ",font=("Arial", 12)).grid(row=0, column=1, sticky="w")
        tk.Label(self.master, text="             ").grid(row=1, column=0, sticky="w")

        # Input Box
        tk.Label(self.master, text="Berat Q").grid(row=1, column=1, sticky="w")
        self.Q_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6,font=("Arial", 12))
        self.Q_entry.grid(row=1, column=2)
        tk.Label(self.master, text="(kN)").grid(row=1, column=3, padx=0, pady=0, sticky="w")

        tk.Label(self.master, text="Momen Sumbu X").grid(row=2, column=1, sticky="w")
        self.Mx_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6,font=("Arial", 12))
        self.Mx_entry.grid(row=2, column=2)
        tk.Label(self.master, text="(kNm)").grid(row=2, column=3, padx=0, pady=0, sticky="w")

        tk.Label(self.master, text="Momen Sumbu Y").grid(row=3, column=1, sticky="w")
        self.My_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6,font=("Arial", 12))
        self.My_entry.grid(row=3, column=2)
        tk.Label(self.master, text="(kNm)").grid(row=3, column=3, padx=0, pady=0, sticky="w")


        tk.Label(self.master, text="Lebar pondasi B").grid(row=4, column=1, sticky="w")
        self.B_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6,font=("Arial", 12))
        self.B_entry.grid(row=4, column=2)
        tk.Label(self.master, text="(m)").grid(row=4, column=3, padx=0, pady=0, sticky="w")


        tk.Label(self.master, text="Panjang pondasi L").grid(row=5, column=1, sticky="w")
        self.L_entry = tk.Entry(self.master, validate="key", validatecommand=(root.register(validate_input), '%S'),width=6, font=("Arial", 12))
        self.L_entry.grid(row=5, column=2)
        tk.Label(self.master, text="(m)").grid(row=5, column=3, padx=0, pady=0, sticky="w")



        """sher = tk.IntVar()
        Radiobutton(self.master, text="General Shear", variable=sher, value=0).grid(row=1, column=4)
        Radiobutton(self.master, text="Local Shear", variable=sher, value=1).grid(row=2, column=4)
"""

        # Calculate Button
        tk.Button(self.master, text="Hitung Tegangan", command=self.calculate).grid(row=10, column=4, rowspan=2, columnspan=2,  padx=40, pady=25)
        # Info Author Button
        tk.Button(self.master, text="Info", command=self.infobruh).grid(row=10, column=1,  padx=0, pady=15, sticky=tk.W)

    def infobruh(self):
        tk.messagebox.showinfo("Info", "Program ini dibuat oleh :\n\n   1. Ridho Rizky Febriansyah 2035221071 (Coder)\n   2. Izzati Eka Nisa 2035221047 (Team)\n   3. Nadia Aura Deka 2035221050 (Team)\n   4. Fiza Eka Salsabilah 2035221084 (Team)\n   5. Salsabilla shafi A 2035221089 (Team)\n   6. Affarel Sebatian N 2035221056 (Team)\n   7. Abdillah Pandega W. 2035221065 (Team)\n   8. Samudro Luhur Trapsilo 2035221086 (Team)\n   9. Syarifuddin Putra P 2035221053 (Team)\n   10. I Made Raghayana 2035221057 (Team)\n   11. Barka Hilman P. 2035221051 (Team)\n")
   


    def calculate(self):
        # Get input values
        try:
            Q = float(self.Q_entry.get())
            Mx = float(self.Mx_entry.get())
            My = float(self.My_entry.get())
            B = float(self.B_entry.get())
            L = float(self.L_entry.get())
            
        # Jendela eror saat vaule salah
        except Exception as e:
            messagebox.showerror("Error", "Tolong Input data dengan benar")


        # perhitungan
        A = B * L
        Xmax = B/2
        Ymax = L/2
        Ix = (1/12)*(B*pow(L,3))
        Iy = (1/12)*(L*pow(B,3))
        qmax = Q/A + (Mx*Ymax)/Ix + (My*Xmax)/Iy
        qmin = Q/A - (Mx*Ymax)/Ix - (My*Xmax)/Iy

        # Simpan Hasil ke dalam File Word
        # import modul python-docx

        # membuat objek document
        doc = docx.Document()


        # Add a Title to the document 
        doc.add_heading('Tegangan Di Dasar Pondasi Dangkal', 0)


        # menambahkan teks ke dokumem
        para = doc.add_paragraph('Diketahui : ').add_run()
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        table = doc.add_table(rows=2, cols=5)

        # mengisi sel-sel tabel dengan teks
        table.cell(0, 0).text = "Q"
        table.cell(0, 1).text = "Mx"
        table.cell(0, 2).text = "My"
        table.cell(0, 3).text = "B"
        table.cell(0, 4).text = "L"
        table.cell(1, 0).text = str(Q)+' (kN)'
        table.cell(1, 1).text = str(Mx)+' (kNm)'
        table.cell(1, 2).text = str(My)+' (kNm)'
        table.cell(1, 3).text = str(B)+' (m)'
        table.cell(1, 4).text = str(L)+' (m)'


        # menambahkan style tabel baru
        new_style = doc.styles.add_style('MyTableStyle', docx.enum.style.WD_STYLE_TYPE.TABLE)
        new_style.base_style = doc.styles['Table Grid']  # gaya dasar tabel
        new_style.font.name = 'Times New Roman'  # jenis huruf
        new_style.font.size = docx.shared.Pt(12)  # ukuran huruf
        new_style.paragraph_format.alignment = docx.enum.text.WD_TAB_ALIGNMENT.CENTER  # alignment teks dalam sel
        new_style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'Arial')  # jenis huruf Asia Timur
        # menerapkan gaya baru pada tabel
        table.style = 'MyTableStyle'


        para = doc.add_paragraph().add_run('\nDitanya :')
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        para = doc.add_paragraph().add_run('Hitung Ix, Iy, Xmax, Ymax, qmax, dan qmin.\n')
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)


        para = doc.add_paragraph().add_run('Dijawab :')
        para.font.name = 'Times New Roman'
        para.font.size = Pt(12)

        # menyisipkan gambar ke dalam dokumen
        # doc.add_picture('assets\ilus1.png'), width=Cm(2), height=Cm(2))

        para = doc.add_paragraph()
        para.add_run("A	= B x L\n")
        para.add_run(" 	= "+str(B)+" x "+str(L)+"\n")
        para.add_run(" 	= "+str(A)+" m2\n")
        para.add_run("Ix	= (1/12) x (B x L^3)\n")
        para.add_run(" 	= (1/12) x ("+str(B)+" x "+str(pow(L,3))+")\n")
        para.add_run(" 	= "+str(Ix)+" m4\n")
        para.add_run("Iy	= (1/12) x (L x B^3)\n")
        para.add_run("	= (1/12) x ("+str(L)+" x "+str(pow(B,3))+")\n")
        para.add_run("	= "+str(Iy)+" m4\n")
        para.add_run("Xmax	= B / 2\n")
        para.add_run("	= "+str(B)+" / 2 \n")
        para.add_run("	= "+str(Xmax)+" m\n")
        para.add_run("Ymax	= L / 2\n")
        para.add_run("	= "+str(L)+" / 2 \n")
        para.add_run("	= "+str(Ymax)+" m\n")

        para.add_run("qmax	= Q/A + (Mx · Ymax)/Ix + (My · Xmax)/Iy\n")
        para.add_run("	= "+str(Q)+"/"+str(A)+" + ("+str(Mx)+" x "+str(Ymax)+")/"+str(Ix)+" + ("+str(My)+" x "+str(Xmax)+")/"+str(Iy)+"\n")
        para.add_run("	= "+str(qmax)+" kN/m2\n")

        para.add_run("qmin	= Q/A - (Mx x Ymax)/Ix - (My x Xmax)/Iy\n")        
        para.add_run("	= "+str(Q)+"/"+str(A)+" - ("+str(Mx)+" x "+str(Ymax)+")/"+str(Ix)+" - ("+str(My)+" x "+str(Xmax)+")/"+str(Iy)+"\n")
        para.add_run("	= "+str(qmin)+" kN/m2\n")

        table = doc.add_table(rows=2, cols=6)

        # mengisi sel-sel tabel dengan teks
        table.cell(0, 0).text = "Ix" 
        table.cell(0, 1).text = "Iy"
        table.cell(0, 2).text = "Xmax"
        table.cell(0, 3).text = "Ymax"
        table.cell(0, 4).text = "qmax"
        table.cell(0, 5).text = "qmin"
        table.cell(1, 0).text = str(Ix)+' (m4)'
        table.cell(1, 1).text = str(Iy)+' (m4)'
        table.cell(1, 2).text = str(Xmax)+' (m)'
        table.cell(1, 3).text = str(Ymax)+' (m)'
        table.cell(1, 4).text = str(qmax)+' (kN/m2)'
        table.cell(1, 5).text = str(qmin)+' (kN/m2)'
        table.style = 'MyTableStyle'


        # menyimpan dokumen ke dalam file
        doc.save('Hasil_Perhitungan_Tegangan_Dasar_Pondasi.docx')


        """# Simpan Hasil ke dalam File Excel
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.title = "Hasil Perhitungan Teg"
        sheet['A1'] = "Tegangan Di Dasar Pondasi Dangkal"
        l1 = ["Ix","Iy","Xmax","Ymax","qmax","qmin"]
        sheet.append(l1)
        l2 = [Ix,Iy,Xmax,Ymax,qmax,qmin,(Mx*Ymax)/Ix,(My*Xmax)/Iy]
        sheet.append(l2)
        
        # Save file Excel
        wb.save("hasil_perhitungan_teg.xlsx")"""

        # Tampilkan pesan berhasil
        tk.messagebox.showinfo("Sukses", "Perhitungan berhasil. Hasil perhitungan tersimpan di Hasil_Perhitungan_Tegangan_Dasar_Pondasi.docx")

if __name__ == "__main__":
    root = tk.Tk()
    root.iconbitmap(ico_path)
    root.geometry("640x200")
    root.resizable(False, False)
    app = App(master=root)
    app.mainloop()

